from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

import docx
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

from catalog.config import get_settings
from catalog.documents.extract import extract_text
from catalog.documents.ingest import allocate_rel_path, safe_filename
from catalog.storage.db import Database
from catalog.storage.repo_document import get_document
from catalog.storage.repo_session_document import list_session_documents

EXPORT_DIR = "export"

_FRONTMATTER_RE = re.compile(r"\A---[ \t]*\n(.*?\n)---[ \t]*(?:\n|\Z)", re.DOTALL)
_WIKI_RE = re.compile(r"(?<!!)\[\[([^\]]+)\]\]")
_HR_RE = re.compile(r"-{3,}")
_UL_RE = re.compile(r"^[ \t]*[-*+][ \t]+(.*)$")
_OL_RE = re.compile(r"^[ \t]*\d{1,9}[.)][ \t]+(.*)$")
_SEP_CELL_RE = re.compile(r":?-{3,}:?")
_LINKS_HEADING = "Ссылки"
_LINKS_EXPORT_HEADING = "Источники"
_MONO_FONT = "Courier New"


def render_docx(md: str, *, template: Path | None = None) -> bytes:
    return render_docx_parts([md], template=template)


def render_docx_parts(
    parts: list[str],
    *,
    template: Path | None = None,
    title: str = "",
) -> bytes:
    document = _open_document(template)
    if title.strip():
        document.core_properties.title = title.strip()
    for index, md in enumerate(parts):
        if index > 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        body, props = _split_frontmatter(md or "")
        if index == 0 and not title.strip():
            if props.get("title"):
                document.core_properties.title = props["title"]
            if props.get("author"):
                document.core_properties.author = props["author"]
        elif index == 0 and props.get("author"):
            document.core_properties.author = props["author"]
        _render_blocks(document, _rewrite_wikilinks(body))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def count_md_structure(md: str) -> tuple[int, int]:
    body, _props = _split_frontmatter(md or "")
    body = _rewrite_wikilinks(body)
    headings = 0
    table_rows = 0
    lines = body.split("\n")
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                index += 1
            if index < len(lines):
                index += 1
            continue
        if _parse_heading(line) is not None:
            headings += 1
            index += 1
            continue
        stripped = line.lstrip()
        if stripped.startswith("|") and "|" in stripped[1:]:
            while index < len(lines):
                candidate = lines[index].lstrip()
                if not (candidate.startswith("|") and "|" in candidate[1:]):
                    break
                cells = _split_table_row(candidate)
                if not _is_separator_row(cells):
                    table_rows += 1
                index += 1
            continue
        index += 1
    return headings, table_rows


def count_extracted_table_rows(text: str) -> int:
    rows = 0
    for line in text.split("\n"):
        stripped = line.lstrip()
        if not (stripped.startswith("|") and "|" in stripped[1:]):
            continue
        if not _is_separator_row(_split_table_row(stripped)):
            rows += 1
    return rows


def count_docx_headings(path: str | Path) -> int:
    document = docx.Document(str(path))
    count = 0
    for paragraph in document.paragraphs:
        style = paragraph.style
        if style is None:
            continue
        name = style.name or ""
        if name.startswith("Heading"):
            count += 1
    return count


def _resolve_template(workspace: Path, template: str) -> Path | None:
    explicit = (template or "").strip()
    if explicit:
        path = Path(explicit)
        if not path.is_absolute():
            path = workspace / path
        if not path.is_file():
            raise FileNotFoundError(str(path))
        return path
    raw = (get_settings().docx_template or "").strip()
    if not raw:
        return None
    path = Path(raw)
    if not path.is_absolute():
        path = workspace / path
    if path.is_file():
        return path
    return None


def write_export_docx(
    db: Database,
    workspace_dir: str | Path,
    doc_ids: list[str],
    *,
    title: str = "",
    template: str = "",
    session_id: str | None = None,
) -> dict[str, Any]:
    if session_id is not None:
        attached_ids = {row.id for row in list_session_documents(db, session_id)}
        for doc_id in doc_ids:
            if doc_id not in attached_ids:
                return {"error": "document_not_available_in_session"}

    if not doc_ids:
        return {"error": "doc_ids required"}

    workspace = Path(workspace_dir)
    parts: list[str] = []
    first_title = ""
    expected_headings = 0
    expected_tables = 0
    for doc_id in doc_ids:
        row = get_document(db, doc_id)
        if row is None:
            return {"error": "document not found"}
        if not first_title:
            first_title = row.title
        text = extract_text(str(workspace / row.path), row.kind)
        parts.append(text)
        headings, tables = count_md_structure(text)
        expected_headings += headings
        expected_tables += tables

    try:
        template_path = _resolve_template(workspace, template)
    except FileNotFoundError:
        return {"error": "template not found"}

    export_title = (title or "").strip() or first_title or "export"
    rel_path = allocate_rel_path(
        workspace, safe_filename(export_title, ".docx"), subdir=EXPORT_DIR
    )
    dest = workspace / rel_path
    dest.write_bytes(
        render_docx_parts(parts, template=template_path, title=export_title)
    )

    extracted = extract_text(str(dest), "docx")
    actual_headings = count_docx_headings(dest)
    actual_tables = count_extracted_table_rows(extracted)
    ok = actual_headings == expected_headings and actual_tables == expected_tables
    return {
        "ok": ok,
        "path": rel_path,
        "headings": actual_headings,
        "tables": actual_tables,
    }


def _open_document(template: Path | None) -> docx.Document:
    if template is None:
        return docx.Document()
    document = docx.Document(str(template))
    body = document.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def _split_frontmatter(text: str) -> tuple[str, dict[str, str]]:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if text.startswith("\ufeff"):
        text = text[1:]
    match = _FRONTMATTER_RE.match(text)
    if match is None:
        return text, {}
    props: dict[str, str] = {}
    for line in match.group(1).splitlines():
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        key = key.strip().lower()
        value = value.strip().strip("\"'")
        if key in ("title", "author") and value:
            props[key] = value
    return text[match.end() :], props


def _rewrite_wikilinks(text: str) -> str:
    parts: list[str] = []
    index = 0
    length = len(text)
    while index < length:
        fence = text.find("```", index)
        if fence == -1:
            parts.append(_WIKI_RE.sub(_wikilink_replace, text[index:]))
            break
        parts.append(_WIKI_RE.sub(_wikilink_replace, text[index:fence]))
        closing = text.find("```", fence + 3)
        if closing == -1:
            parts.append(text[fence:])
            break
        parts.append(text[fence : closing + 3])
        index = closing + 3
    return "".join(parts)


def _wikilink_replace(match: re.Match[str]) -> str:
    inner = match.group(1)
    alias_sep = inner.find("|")
    if alias_sep >= 0:
        return inner[alias_sep + 1 :].strip()
    target = inner
    heading_sep = target.find("#")
    if heading_sep >= 0:
        target = target[:heading_sep]
    return target.strip()


def _render_blocks(document: docx.Document, text: str) -> None:
    lines = text.split("\n")
    index = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        _add_paragraph(document, " ".join(line.strip() for line in paragraph_lines))
        paragraph_lines.clear()

    while index < len(lines):
        line = lines[index]
        if line.strip() == "":
            flush_paragraph()
            index += 1
            continue
        if line.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            _add_code(document, "\n".join(code_lines).rstrip("\n"))
            continue
        heading = _parse_heading(line)
        if heading is not None:
            flush_paragraph()
            level, heading_text = heading
            if heading_text == _LINKS_HEADING:
                heading_text = _LINKS_EXPORT_HEADING
            _add_heading(document, level, heading_text)
            index += 1
            continue
        if _HR_RE.fullmatch(line.strip()):
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue
        stripped = line.lstrip()
        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            rows: list[list[str]] = []
            while index < len(lines):
                candidate = lines[index].lstrip()
                if not (candidate.startswith("|") and "|" in candidate[1:]):
                    break
                rows.append(_split_table_row(candidate))
                index += 1
            if len(rows) >= 2 and _is_separator_row(rows[1]):
                del rows[1]
            if rows:
                _add_table(document, rows)
            continue
        unordered = _UL_RE.match(line)
        if unordered is not None:
            flush_paragraph()
            while index < len(lines):
                item = _UL_RE.match(lines[index])
                if item is None:
                    break
                _add_list_item(document, item.group(1), ordered=False)
                index += 1
            continue
        ordered = _OL_RE.match(line)
        if ordered is not None:
            flush_paragraph()
            while index < len(lines):
                item = _OL_RE.match(lines[index])
                if item is None:
                    break
                _add_list_item(document, item.group(1), ordered=True)
                index += 1
            continue
        paragraph_lines.append(line)
        index += 1
    flush_paragraph()


def _parse_heading(line: str) -> tuple[int, str] | None:
    if not line.startswith("#"):
        return None
    level = 0
    while level < len(line) and level < 6 and line[level] == "#":
        level += 1
    if level < len(line) and line[level] not in " \t":
        return None
    text = line[level:].strip()
    if text.endswith("#"):
        text = text.rstrip("#").rstrip()
    return level, text


def _split_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|") and not raw.endswith("\\|"):
        raw = raw[:-1]
    cells: list[str] = []
    current: list[str] = []
    index = 0
    while index < len(raw):
        if raw[index] == "\\" and index + 1 < len(raw) and raw[index + 1] == "|":
            current.append("|")
            index += 2
            continue
        if raw[index] == "|":
            cells.append("".join(current).strip())
            current = []
            index += 1
            continue
        current.append(raw[index])
        index += 1
    cells.append("".join(current).strip())
    return cells


def _is_separator_row(cells: list[str]) -> bool:
    if not cells:
        return False
    return all(_SEP_CELL_RE.fullmatch(cell.replace(" ", "")) is not None for cell in cells)


def _add_heading(document: docx.Document, level: int, text: str) -> None:
    try:
        paragraph = document.add_heading("", level=level)
    except ValueError:
        paragraph = document.add_paragraph()
    _add_runs(paragraph, text)


def _add_paragraph(document: docx.Document, text: str) -> None:
    paragraph = document.add_paragraph()
    _add_runs(paragraph, text)


def _add_list_item(document: docx.Document, text: str, *, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    try:
        paragraph = document.add_paragraph("", style=style)
    except ValueError:
        paragraph = document.add_paragraph()
    _add_runs(paragraph, text)


def _add_code(document: docx.Document, text: str) -> None:
    if text == "":
        paragraph = document.add_paragraph()
        _apply_monospace(paragraph.add_run(""))
        return
    for line in text.split("\n"):
        paragraph = document.add_paragraph()
        _apply_monospace(paragraph.add_run(line))


def _add_table(document: docx.Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    if width == 0:
        return
    try:
        table = document.add_table(rows=len(rows), cols=width, style="Table Grid")
    except ValueError:
        table = document.add_table(rows=len(rows), cols=width)
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.text = ""
            _add_runs(cell.paragraphs[0], value)


def _add_runs(paragraph: object, text: str) -> None:
    for chunk, bold, italic, code in _iter_inlines(text):
        run = paragraph.add_run(chunk)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            _apply_monospace(run)


def _apply_monospace(run: object) -> None:
    run.font.name = _MONO_FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), _MONO_FONT)
    r_fonts.set(qn("w:hAnsi"), _MONO_FONT)
    r_fonts.set(qn("w:eastAsia"), _MONO_FONT)
    r_fonts.set(qn("w:cs"), _MONO_FONT)


def _iter_inlines(text: str) -> list[tuple[str, bool, bool, bool]]:
    runs: list[tuple[str, bool, bool, bool]] = []
    buffer: list[str] = []
    index = 0
    length = len(text)

    def flush() -> None:
        if buffer:
            runs.append(("".join(buffer), False, False, False))
            buffer.clear()

    while index < length:
        if text[index] == "`":
            closing = text.find("`", index + 1)
            if closing != -1:
                flush()
                runs.append((text[index + 1 : closing], False, False, True))
                index = closing + 1
                continue
        if text.startswith("**", index):
            closing = text.find("**", index + 2)
            if closing != -1:
                flush()
                for chunk, _bold, italic, code in _iter_inlines(text[index + 2 : closing]):
                    runs.append((chunk, not code, italic, code))
                index = closing + 2
                continue
        if text[index] == "*":
            closing = index + 1
            while closing < length:
                if text[closing] == "*" and not text.startswith("**", closing):
                    break
                closing += 1
            if closing < length and text[closing] == "*":
                flush()
                for chunk, bold, _italic, code in _iter_inlines(text[index + 1 : closing]):
                    runs.append((chunk, bold, not code, code))
                index = closing + 1
                continue
        buffer.append(text[index])
        index += 1
    flush()
    return runs
