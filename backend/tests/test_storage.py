from __future__ import annotations

import asyncio
import sqlite3
from pathlib import Path

import pytest

from catalog.documents.extract import extract_text
from catalog.documents.ingest import ingest_file
from catalog.documents.tools import build_document_tools
from catalog.storage.db import Database
from catalog.agent.trace import Trace
from catalog.skills.repo_run import create_run, finish_run, get_run
from catalog.storage.repo_document import (
    DocumentRow,
    create_document,
    delete_document,
    get_document,
    list_documents,
    list_documents_by_kind,
    reconcile_orphans,
)
from catalog.storage.repo_session import create_session
from catalog.storage.repo_session_document import (
    attach_documents,
    detach_documents,
    list_session_documents,
)



@pytest.fixture()
def db() -> Database:
    d = Database(":memory:")
    d.init_schema()
    return d


def _table_names(db: Database) -> set[str]:
    with db.connect() as conn:
        rows = conn.execute("SELECT name FROM sqlite_master WHERE type = 'table'").fetchall()
    return {r["name"] for r in rows}


def test_schema_creates_all_tables(db: Database) -> None:
    assert {
        "document",
        "session",
        "session_document",
        "session_skill",
        "message",
        "skill",
        "skill_run",
        "session_artifact",
        "custom_check",
    } <= _table_names(db)


def test_init_schema_is_idempotent(db: Database) -> None:
    db.init_schema()  # second run must not raise
    assert {
        "document",
        "session",
        "session_document",
        "session_skill",
        "message",
        "skill",
        "skill_run",
        "session_artifact",
        "custom_check",
    } <= _table_names(db)


def test_workspace_user_version_set(db: Database) -> None:
    from catalog.storage.schema import WORKSPACE_USER_VERSION

    assert db.user_version() == WORKSPACE_USER_VERSION


def test_app_schema_tables() -> None:
    from catalog.storage.schema import (
        APP_ADDITIVE_MIGRATIONS,
        APP_SCHEMA,
        APP_USER_VERSION,
    )

    d = Database(":memory:")
    d.init_schema(APP_SCHEMA, APP_USER_VERSION, migrations=APP_ADDITIVE_MIGRATIONS)
    assert {"workspace_registry", "app_settings"} <= _table_names(d)
    assert d.user_version() == APP_USER_VERSION
    with d.connect() as conn:
        cols = {
            row["name"]
            for row in conn.execute("PRAGMA table_info(app_settings)").fetchall()
        }
    assert {"openrouter_api_key", "zai_api_key"} <= cols


def test_create_and_get_document(db: Database) -> None:
    row = create_document(db, title="T", path="documents/x.md", kind="md")
    assert isinstance(row, DocumentRow)
    assert len(row.id) == 32  # uuid4 hex
    assert row.title == "T"

    fetched = get_document(db, row.id)
    assert fetched is not None
    assert fetched.id == row.id
    assert fetched.kind == "md"
    assert fetched.path == "documents/x.md"

    assert get_document(db, "missing") is None


def test_list_documents_and_by_kind(db: Database) -> None:
    create_document(db, title="a", path="documents/a.md", kind="md")
    create_document(db, title="b", path="documents/b.md", kind="md")
    create_document(db, title="c", path="documents/c.docx", kind="docx")

    assert len(list_documents(db)) == 3
    assert len(list_documents_by_kind(db, "md")) == 2
    assert len(list_documents_by_kind(db, "docx")) == 1
    assert list_documents_by_kind(db, "result_md") == []


def test_attach_documents_idempotent(db: Database) -> None:
    session_id = create_session(db)
    a = create_document(db, title="Alpha", path="documents/a.md", kind="md")
    b = create_document(db, title="Beta", path="documents/b.md", kind="md")

    assert attach_documents(db, session_id, [a.id, b.id]) == []
    assert attach_documents(db, session_id, [a.id, b.id]) == []
    assert attach_documents(db, session_id, ["missing-id"]) == ["missing-id"]
    assert attach_documents(db, session_id, [a.id, "missing-id", "missing-id"]) == [
        "missing-id"
    ]
    assert attach_documents(db, session_id, []) == []

    docs = list_session_documents(db, session_id)
    assert [d.id for d in docs] == [a.id, b.id]
    assert [d.title for d in docs] == ["Alpha", "Beta"]


def test_detach_documents_idempotent(db: Database) -> None:
    session_id = create_session(db)
    a = create_document(db, title="Alpha", path="documents/a.md", kind="md")
    b = create_document(db, title="Beta", path="documents/b.md", kind="md")
    attach_documents(db, session_id, [a.id, b.id])

    assert detach_documents(db, session_id, [a.id]) == 1
    assert [d.id for d in list_session_documents(db, session_id)] == [b.id]
    assert detach_documents(db, session_id, [a.id]) == 0
    assert detach_documents(db, session_id, []) == 0
    assert get_document(db, a.id) is not None
    assert [d.id for d in list_session_documents(db, session_id)] == [b.id]


def test_ingest_md_and_read(db: Database, tmp_path: Path) -> None:
    row = ingest_file(db, tmp_path, filename="note.md", content=b"# Hello\nworld")
    assert row.kind == "md"
    assert row.title == "note"
    assert row.path == "note.md"
    assert (tmp_path / row.path).read_bytes() == b"# Hello\nworld"
    assert row.content_hash is not None
    assert row.size == len(b"# Hello\nworld")

    text = extract_text(str(tmp_path / row.path), row.kind)
    assert "Hello" in text
    assert text == "# Hello\nworld"


def test_ingest_docx_and_read(db: Database, tmp_path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    src = tmp_path / "src.docx"
    doc.save(str(src))
    content = src.read_bytes()

    row = ingest_file(db, tmp_path, filename="report.docx", content=content)
    assert row.kind == "docx"
    assert row.title == "report"
    assert row.path == "report.docx"

    text = extract_text(str(tmp_path / row.path), row.kind)
    assert text == "First paragraph\nSecond paragraph"


def test_unsupported_format_raises(db: Database, tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="unsupported format"):
        ingest_file(db, tmp_path, filename="file.exe", content=b"MZ\x90\x00")


def test_xls_raises_with_hint(db: Database, tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="пересохраните файл как .xlsx"):
        ingest_file(db, tmp_path, filename="old.xls", content=b"not-xlsx")
    assert not (tmp_path / "old.xls").exists()
    assert list_documents(db) == []


def test_ingest_broken_xlsx_does_not_create(db: Database, tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="xlsx"):
        ingest_file(db, tmp_path, filename="broken.xlsx", content=b"not-a-zip")
    assert not (tmp_path / "broken.xlsx").exists()
    assert list_documents(db) == []


def test_ingest_keeps_original_filename(db: Database, tmp_path: Path) -> None:
    row = ingest_file(
        db, tmp_path, filename="Пример Документ.md", content=b"content"
    )
    assert row.title == "Пример Документ"
    assert row.path == "Пример Документ.md"
    assert (tmp_path / row.path).read_bytes() == b"content"


def test_ingest_collision_adds_suffix(db: Database, tmp_path: Path) -> None:
    first = ingest_file(db, tmp_path, filename="note.md", content=b"a")
    second = ingest_file(db, tmp_path, filename="note.md", content=b"b")
    assert first.path == "note.md"
    assert second.path == "note-1.md"
    assert (tmp_path / second.path).read_bytes() == b"b"


def test_ingest_garbage_filename_kept(db: Database, tmp_path: Path) -> None:
    row = ingest_file(db, tmp_path, filename="@@@!!!.md", content=b"content")
    assert row.title == "@@@!!!"
    assert row.path == "@@@!!!.md"
    assert (tmp_path / row.path).read_bytes() == b"content"


def test_ingest_docx_extension_still_validated(db: Database, tmp_path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("Текст")
    src = tmp_path / "src.docx"
    doc.save(str(src))

    row = ingest_file(
        db, tmp_path, filename="Годовой отчёт.docx", content=src.read_bytes()
    )
    assert row.kind == "docx"
    assert row.path == "Годовой отчёт.docx"


def test_read_unknown_doc_error(db: Database, tmp_path: Path) -> None:
    reg = build_document_tools(db, tmp_path)
    entry = reg.get("read_document")
    assert entry is not None
    _, read_fn = entry

    async def _run() -> dict:
        return await read_fn(doc_id="nope")

    assert asyncio.run(_run()) == {"error": "document not found"}


def test_list_documents_tool(db: Database, tmp_path: Path) -> None:
    ingest_file(db, tmp_path, filename="a.md", content=b"a")
    ingest_file(db, tmp_path, filename="b.md", content=b"b")
    reg = build_document_tools(db, tmp_path)
    entry = reg.get("list_documents")
    assert entry is not None
    _, list_fn = entry

    async def _run() -> list[dict]:
        return await list_fn()

    items = asyncio.run(_run())
    assert len(items) == 2
    assert {it["kind"] for it in items} == {"md"}
    assert all({"id", "title", "kind"} == set(it) for it in items)


def test_read_document_tool_md(db: Database, tmp_path: Path) -> None:
    row = ingest_file(db, tmp_path, filename="note.md", content=b"body text")
    reg = build_document_tools(db, tmp_path)
    entry = reg.get("read_document")
    assert entry is not None
    _, read_fn = entry

    async def _run() -> dict:
        return await read_fn(doc_id=row.id)

    result = asyncio.run(_run())
    assert result == {"title": "note", "kind": "md", "text": "body text"}


def test_document_tools_registered(db: Database, tmp_path: Path) -> None:
    reg = build_document_tools(db, tmp_path)
    assert reg.names() == ["list_documents", "read_document"]

    specs = {s.name: s for s in reg.specs()}
    assert specs["list_documents"].parameters == {"type": "object", "properties": {}}
    assert specs["read_document"].parameters == {
        "type": "object",
        "properties": {"doc_id": {"type": "string"}},
        "required": ["doc_id"],
    }


def test_row_factory_is_set(db: Database) -> None:
    # connect() must yield connections whose rows are sqlite3.Row.
    with db.connect() as conn:
        assert conn.row_factory is sqlite3.Row


def test_delete_document_removes_file_and_row(db: Database, tmp_path: Path) -> None:
    row = ingest_file(db, tmp_path, filename="note.md", content=b"body")
    assert (tmp_path / row.path).is_file()

    deleted = delete_document(db, tmp_path, row.id)
    assert deleted is not None
    assert deleted.id == row.id
    assert not (tmp_path / row.path).exists()
    assert get_document(db, row.id) is None
    assert delete_document(db, tmp_path, row.id) is None


def test_reconcile_orphans_removes_missing_files(db: Database, tmp_path: Path) -> None:
    kept = ingest_file(db, tmp_path, filename="keep.md", content=b"keep")
    orphan = ingest_file(db, tmp_path, filename="gone.md", content=b"gone")
    (tmp_path / orphan.path).unlink()

    removed = reconcile_orphans(db, tmp_path)
    assert removed == [orphan.id]
    assert get_document(db, orphan.id) is None
    assert get_document(db, kept.id) is not None
    assert (tmp_path / kept.path).is_file()


def test_document_schema_has_extracted_text_column(db: Database) -> None:
    with db.connect() as conn:
        cols = {row["name"] for row in conn.execute("PRAGMA table_info(document)")}
        fts = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'fts%'"
        ).fetchall()
    assert "extracted_text" in cols
    assert "mtime" in cols
    assert "size" in cols
    assert "content_hash" in cols
    assert fts == []


def test_skill_run_schema_has_parent_run_id(db: Database) -> None:
    with db.connect() as conn:
        cols = {row["name"] for row in conn.execute("PRAGMA table_info(skill_run)")}
    assert "parent_run_id" in cols


def test_delete_document_nullifies_skill_run_refs(db: Database, tmp_path: Path) -> None:
    input_a = ingest_file(db, tmp_path, filename="a.md", content=b"a")
    input_b = ingest_file(db, tmp_path, filename="b.md", content=b"b")
    output = ingest_file(db, tmp_path, filename="out.md", content=b"out")
    run_id = create_run(
        db, skill_id="skill1", session_id=None, input_doc_ids=[input_a.id, input_b.id]
    )
    finish_run(
        db,
        run_id,
        status="ok",
        output_doc_id=output.id,
        trace=Trace(),
        result_text="done",
    )

    delete_document(db, tmp_path, input_a.id)
    delete_document(db, tmp_path, output.id)

    run = get_run(db, run_id)
    assert run is not None
    assert run["status"] == "ok"
    assert run["result_text"] == "done"
    assert run["input_doc_id"] == input_b.id
    assert run["input_doc_ids"] == [input_b.id]
    assert run["output_doc_id"] is None


def test_list_documents_tool_reconciles_orphans(db: Database, tmp_path: Path) -> None:
    kept = ingest_file(db, tmp_path, filename="keep.md", content=b"keep")
    orphan = ingest_file(db, tmp_path, filename="gone.md", content=b"gone")
    (tmp_path / orphan.path).unlink()
    reg = build_document_tools(db, tmp_path)
    entry = reg.get("list_documents")
    assert entry is not None
    _, list_fn = entry

    async def _run() -> list[dict]:
        return await list_fn()

    items = asyncio.run(_run())
    assert [it["id"] for it in items] == [kept.id]
    assert get_document(db, orphan.id) is None
