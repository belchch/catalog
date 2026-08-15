from __future__ import annotations

from pathlib import Path

import docx
import pytest
from docx.oxml.ns import qn

from catalog.documents.extract import extract_text

_FIXTURES = Path(__file__).parent / "fixtures"
_SYNTHETIC_DEFECTS = _FIXTURES / "defects_table_synthetic.docx"
_REAL_DEFECTS = _FIXTURES / "defects_table_10_2025-06-09T13_36_18.docx"


def _save(doc: docx.Document, path: Path) -> Path:
    doc.save(str(path))
    return path


def test_extract_docx_table(tmp_path: Path) -> None:
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "city"
    table.cell(1, 0).text = "Anna"
    table.cell(1, 1).text = "Moscow"
    path = _save(doc, tmp_path / "table.docx")

    text = extract_text(str(path), "docx")
    assert "| name | city |" in text
    assert "| --- | --- |" in text
    assert "| Anna | Moscow |" in text


def test_extract_docx_paragraph_table_paragraph(tmp_path: Path) -> None:
    doc = docx.Document()
    doc.add_paragraph("Before")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    doc.add_paragraph("After")
    path = _save(doc, tmp_path / "mix.docx")

    text = extract_text(str(path), "docx")
    assert text.index("Before") < text.index("| H1 | H2 |") < text.index("After")
    assert "| a | b |" in text


def test_extract_docx_merged_cells(tmp_path: Path) -> None:
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Span"
    table.cell(1, 0).text = "L"
    table.cell(1, 1).text = "R"
    path = _save(doc, tmp_path / "merged.docx")

    text = extract_text(str(path), "docx")
    assert "| Span | Span |" in text
    assert "| L | R |" in text


def test_extract_docx_without_tbl_grid(tmp_path: Path) -> None:
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    tbl = table._tbl
    for child in list(tbl):
        if child.tag == qn("w:tblGrid"):
            tbl.remove(child)
    path = _save(doc, tmp_path / "no_grid.docx")

    text = extract_text(str(path), "docx")
    assert "| A | B |" in text
    assert "| C | D |" in text


def test_extract_docx_nested_table_after_parent(tmp_path: Path) -> None:
    doc = docx.Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.text = "parent"
    inner = cell.add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "ih1"
    inner.cell(0, 1).text = "ih2"
    inner.cell(1, 0).text = "n1"
    inner.cell(1, 1).text = "n2"
    path = _save(doc, tmp_path / "nested.docx")

    text = extract_text(str(path), "docx")
    assert "| parent |" in text
    assert "| ih1 | ih2 |" in text
    assert "| n1 | n2 |" in text
    assert text.index("| parent |") < text.index("| ih1 | ih2 |")


def test_extract_docx_escapes_pipes_and_newlines(tmp_path: Path) -> None:
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "a|b"
    table.cell(0, 1).text = "line1\nline2"
    table.cell(1, 0).text = "ok"
    table.cell(1, 1).text = "c\r\nd"
    path = _save(doc, tmp_path / "special.docx")

    text = extract_text(str(path), "docx")
    assert "| a\\|b | line1 line2 |" in text
    assert "| ok | c  d |" in text
    assert "| a|b |" not in text


def test_extract_docx_without_tables_unchanged(tmp_path: Path) -> None:
    doc = docx.Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    path = _save(doc, tmp_path / "plain.docx")

    text = extract_text(str(path), "docx")
    assert text == "First paragraph\nSecond paragraph"


def test_extract_docx_synthetic_defects_table() -> None:
    text = extract_text(str(_SYNTHETIC_DEFECTS), "docx")
    assert len(text) >= 4000
    assert "Таблица недостатков" in text
    assert "ГОСТ 31173-2016" in text
    assert text.index("Таблица недостатков") < text.index("|")


@pytest.mark.skipif(not _REAL_DEFECTS.exists(), reason="real defects fixture not present")
def test_extract_docx_real_defects_table() -> None:
    text = extract_text(str(_REAL_DEFECTS), "docx")
    assert len(text) >= 4000
    assert "ГОСТ 31173-2016" in text
