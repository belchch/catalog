from __future__ import annotations

from io import BytesIO
from pathlib import Path

import docx
from docx.oxml.ns import qn

from catalog.documents.export_docx import render_docx
from catalog.documents.extract import extract_text


def _document(data: bytes) -> docx.Document:
    return docx.Document(BytesIO(data))


def _extract(data: bytes, tmp_path: Path) -> str:
    path = tmp_path / "roundtrip.docx"
    path.write_bytes(data)
    return extract_text(str(path), "docx")


def _style_names(document: docx.Document) -> list[str]:
    return [paragraph.style.name for paragraph in document.paragraphs if paragraph.text]


def _has_page_break(paragraph: object) -> bool:
    return bool(paragraph._p.findall(".//" + qn("w:br")))


def test_render_docx_returns_openable_bytes() -> None:
    data = render_docx("# Title\n\nHello")
    assert data.startswith(b"PK")
    document = _document(data)
    assert [p.text for p in document.paragraphs] == ["Title", "Hello"]


def test_render_docx_heading_levels() -> None:
    md = "\n".join(f"{'#' * level} H{level}" for level in range(1, 7))
    document = _document(render_docx(md))
    assert _style_names(document) == [f"Heading {level}" for level in range(1, 7)]
    assert [p.text for p in document.paragraphs] == [f"H{level}" for level in range(1, 7)]


def test_render_docx_lists() -> None:
    md = "- bullet\n- two\n\n1. one\n2. two"
    document = _document(render_docx(md))
    styles = _style_names(document)
    assert styles[:2] == ["List Bullet", "List Bullet"]
    assert styles[2:] == ["List Number", "List Number"]
    assert [p.text for p in document.paragraphs] == ["bullet", "two", "one", "two"]


def test_render_docx_table_unescapes_pipe(tmp_path: Path) -> None:
    md = "| a | b |\n| --- | --- |\n| x\\|y | z |"
    data = render_docx(md)
    document = _document(data)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "a"
    assert table.cell(1, 0).text == "x|y"
    assert table.cell(1, 1).text == "z"
    extracted = _extract(data, tmp_path)
    assert "| a | b |" in extracted
    assert "| x\\|y | z |" in extracted


def test_render_docx_bold_italic_and_inline_code() -> None:
    document = _document(render_docx("This is **bold** and *italic* and `code`."))
    runs = list(document.paragraphs[0].runs)
    by_text = {run.text: run for run in runs}
    assert by_text["bold"].bold
    assert not by_text["bold"].italic
    assert by_text["italic"].italic
    assert not by_text["italic"].bold
    assert by_text["code"].font.name == "Courier New"
    assert document.paragraphs[0].text == "This is bold and italic and code."


def test_render_docx_wikilink_forms() -> None:
    md = "See [[Target|alias]] and [[Target]] and [[Target#Heading]]."
    document = _document(render_docx(md))
    assert document.paragraphs[0].text == "See alias and Target and Target."


def test_render_docx_keeps_embed_wikilink() -> None:
    document = _document(render_docx("Keep ![[embed]] as is."))
    assert document.paragraphs[0].text == "Keep ![[embed]] as is."


def test_render_docx_frontmatter_to_core_properties() -> None:
    md = "---\ntitle: Report\nauthor: Ivan\n---\n\n# Body\n"
    data = render_docx(md)
    document = _document(data)
    assert document.core_properties.title == "Report"
    assert document.core_properties.author == "Ivan"
    texts = [p.text for p in document.paragraphs]
    assert texts == ["Body"]
    assert all("---" not in text and "title:" not in text for text in texts)


def test_render_docx_links_section_becomes_sources() -> None:
    md = "body\n\n## Ссылки\n\n- [[cover-letter-ea411722]]\n"
    document = _document(render_docx(md))
    texts = [p.text for p in document.paragraphs]
    assert "Ссылки" not in texts
    assert "Источники" in texts
    assert "cover-letter-ea411722" in texts
    heading = next(p for p in document.paragraphs if p.text == "Источники")
    assert heading.style.name == "Heading 2"


def test_render_docx_uses_template_styles(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    template = docx.Document()
    template.styles["Normal"].font.name = "Times New Roman"
    template.add_paragraph("placeholder")
    template.save(str(template_path))

    data = render_docx("Hello from body", template=template_path)
    document = _document(data)
    texts = [p.text for p in document.paragraphs]
    assert "placeholder" not in texts
    assert "Hello from body" in texts
    assert document.styles["Normal"].font.name == "Times New Roman"


def test_render_docx_unsupported_markup_kept() -> None:
    md = "![img](http://x) and <span>html</span> and [link](url)"
    document = _document(render_docx(md))
    assert document.paragraphs[0].text == md


def test_render_docx_page_break_and_fenced_code() -> None:
    md = "Before\n\n---\n\n```\ndef foo():\n    return 1\n```\n"
    document = _document(render_docx(md))
    texts = [p.text for p in document.paragraphs]
    assert "Before" in texts
    assert "def foo():" in texts
    assert "    return 1" in texts
    assert any(_has_page_break(p) for p in document.paragraphs)
    code_run = next(p.runs[0] for p in document.paragraphs if p.text.startswith("def foo"))
    assert code_run.font.name == "Courier New"


def test_render_docx_round_trip_structure(tmp_path: Path) -> None:
    md = (
        "# Title\n\n"
        "Intro text with **bold**.\n\n"
        "## Section\n\n"
        "- item a\n"
        "- item b\n\n"
        "| Col A | Col B |\n"
        "| --- | --- |\n"
        "| 1 | 2 |\n"
        "| 3\\|4 | 5 |\n\n"
        "Closing.\n"
    )
    extracted = _extract(render_docx(md), tmp_path)
    assert extracted.index("Title") < extracted.index("Intro text with bold")
    assert extracted.index("Intro text with bold") < extracted.index("Section")
    assert extracted.index("Section") < extracted.index("item a")
    assert extracted.index("item a") < extracted.index("item b")
    assert extracted.index("item b") < extracted.index("| Col A | Col B |")
    assert "| 1 | 2 |" in extracted
    assert "| 3\\|4 | 5 |" in extracted
    assert extracted.index("| 3\\|4 | 5 |") < extracted.index("Closing.")
